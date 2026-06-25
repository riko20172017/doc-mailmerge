from openpyxl import load_workbook
from mailmerge import MailMerge, MailMergeOptions, OptionAutoUpdateFields, OptionKeepFields
import qrcode
from io import BytesIO
import tempfile
import os
from PIL import Image
from pathlib import Path
import warnings
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy

doc = Document('t.docx')

# Просто все параграфы документа
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if "предметы" in paragraph.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ''

                    # for i, text in enumerate(["111", "222"]):
                    #     new_para = doc.add_paragraph(text)
                    #     parent.insert(index + 1 + i, new_para._element)
                    for item in enumerate(["111", "222"]):
                        cell.add_paragraph(str(item))
doc.save('t1.docx')
